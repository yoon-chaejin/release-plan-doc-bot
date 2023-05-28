import os
import unittest

from docx import Document

from main import release_plan_controller as rpc


class TestReleasePlanController(unittest.TestCase):
    def test_sample(self):
        self.assertEqual(True, True)

    def test_from_docx_file_paths(self):
        # given
        root_dir = os.getcwd()
        file_paths = [
            root_dir + "/resources/sample.docx",
            root_dir + "/resources/sample_copy.docx",
        ]

        # when
        release_plan_controller = rpc.ReleasePlanController.from_docx_file_paths(
            file_paths
        )
        release_plans = release_plan_controller.release_plans

        # then
        self.assertIsInstance(release_plan_controller, rpc.ReleasePlanController)
        self.assertEqual(2, len(release_plans))

    def test_create_daily_report(self):
        # given
        root_dir = os.getcwd()
        file_paths = [
            root_dir + "/resources/sample.docx",
            root_dir + "/resources/sample_copy.docx",
        ]
        release_plan_controller = rpc.ReleasePlanController.from_docx_file_paths(
            file_paths
        )
        release_plans = release_plan_controller.release_plans

        # when
        # create_daily_report 호출
        daily_report = release_plan_controller.create_daily_report()

        # then
        self.assertIsInstance(daily_report, type(Document()))
        self.assertIsNotNone(daily_report.tables[0])

        self.assertEqual(len(daily_report.tables[0].rows), 3)
        self.assertEqual(len(daily_report.tables[0].columns), 2)

        self.assertEqual(daily_report.tables[0].cell(0, 0).text, "작업자")
        self.assertEqual(daily_report.tables[0].cell(0, 1).text, "작업 내용")

        self.assertEqual(
            daily_report.tables[0].cell(1, 0).text,
            release_plans[0].developer_name,
        )
        self.assertEqual(
            daily_report.tables[0].cell(1, 1).text,
            release_plans[0].release_plan_description,
        )
        self.assertEqual(
            daily_report.tables[0].cell(2, 0).text,
            release_plans[1].developer_name,
        )
        self.assertEqual(
            daily_report.tables[0].cell(2, 1).text,
            release_plans[1].release_plan_description,
        )


if __name__ == "__main__":
    unittest.main()
