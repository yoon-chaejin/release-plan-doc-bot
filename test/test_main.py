import unittest
import os
from main.release_plan import ReleasePlanFromDocxFile
from docx import Document
from main.main import create_daily_report

class TestMain(unittest.TestCase):

    def test_sample(self):
        self.assertEqual(True, True)

    def test_create_daily_report(self):
        # given
        root_dir = os.getcwd()
        file_paths = [root_dir + '/resources/sample.docx', root_dir + '/resources/sample_copy.docx']
        release_plans = [ReleasePlanFromDocxFile(i) for i in file_paths]

        # when
        # create_daily_report 호출
        daily_report = create_daily_report(release_plans)
        
        # then
        self.assertIsInstance(daily_report, type(Document()))
        self.assertIsNotNone(daily_report.tables[0])
        
        self.assertEqual(len(daily_report.tables[0].rows), 3)
        self.assertEqual(len(daily_report.tables[0].columns), 2)
        
        self.assertEqual(daily_report.tables[0].cell(0,0).text, "작업자")
        self.assertEqual(daily_report.tables[0].cell(0,1).text, "작업 내용")


if __name__ == '__main__':
    unittest.main()