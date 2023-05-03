from docx import Document

def create_daily_report(release_plans):
    daily_report = Document()

    daily_report.add_table(3,2)
    daily_report.tables[0].cell(0,0).text = "작업자"
    daily_report.tables[0].cell(0,1).text = "작업 내용"

    return daily_report
