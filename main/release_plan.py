from docx import Document


class ReleasePlan:
    def __init__(
        self,
        document=None,
        title=None,
        request_id=None,
        developer_name=None,
        release_plan_description=None,
    ) -> None:
        self.document = document
        self.title = title
        self.request_id = request_id
        self.developer_name = developer_name
        self.release_plan_description = release_plan_description

    @classmethod
    def from_docx_file_path(cls, file_path: str) -> None:
        document = Document(file_path)
        return cls(
            document,
            document.tables[2].cell(0, 1).text,
            document.tables[2].cell(3, 12).text,
            document.tables[2].cell(5, 10).text,
            document.tables[2].cell(10, 1).text,
        )
