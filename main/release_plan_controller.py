from docx import Document

from main import release_plan as rp


class ReleasePlanController:
    def __init__(self, release_plans=None):
        self.release_plans = release_plans

    @classmethod
    def from_docx_file_paths(cls, file_paths: list):
        release_plans = [rp.ReleasePlan.from_docx_file_path(i) for i in file_paths]
        return cls(release_plans)

    def create_daily_report(self):
        daily_report = Document()

        daily_report.add_table(3, 2)
        daily_report.tables[0].cell(0, 0).text = "작업자"
        daily_report.tables[0].cell(0, 1).text = "작업 내용"

        for idx, release_plan in enumerate(self.release_plans):
            daily_report.tables[0].cell(idx + 1, 0).text = release_plan.developer_name
            daily_report.tables[0].cell(
                idx + 1, 1
            ).text = release_plan.release_plan_description

        return daily_report
